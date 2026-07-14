# -*- coding: utf-8 -*-
"""One-page (A4, 12pt) strategy factory summary generator.

Produces docs/전략공장_요약.docx describing the two live books:
  [전략 1] 금리 팩터 포트폴리오 (Sleeve Engine)
  [전략 2] FX 퀀트 전략 공장 (Discovery Factory)
in the same 전략설명/전략내용/리스크관리 format as the reference example.
"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

section = doc.sections[0]
section.page_height = Cm(29.7)
section.page_width = Cm(21.0)
section.top_margin = Cm(1.3)
section.bottom_margin = Cm(1.3)
section.left_margin = Cm(1.7)
section.right_margin = Cm(1.7)

style = doc.styles['Normal']
style.font.name = '맑은 고딕'
style.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
style.font.size = Pt(12)
style.paragraph_format.space_after = Pt(1)
style.paragraph_format.line_spacing = 1.0


def set_font(run, size=12, bold=False):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')


def add_title(text, size=15):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    set_font(r, size=size, bold=True)
    return p


def add_heading(text, size=13):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(text)
    set_font(r, size=size, bold=True)
    return p


def add_body(text, size=12, space_after=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    set_font(r, size=size)
    return p


def add_bullet(text, size=12, indent=0.4, space_after=0):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    set_font(r, size=size)
    return p


def add_subbullet(text, size=12, indent=0.9, space_after=0):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    set_font(r, size=size)
    return p


# ── Title ────────────────────────────────────────────────────────────
add_title("전략 공장 개요: 금리 팩터 포트폴리오 & FX 퀀트 전략 공장")

# ── [전략 1] 금리 팩터 포트폴리오 ───────────────────────────────────
add_heading("[전략 1] 금리 팩터 포트폴리오 (Rates Sleeve Engine)")
add_body(
    "미국 국채선물 자산군을 대상으로, 서로 독립적인 4개 팩터(추세·가치·캐리·정책)를 "
    "결합하여 지속적으로 리스크가 조절된 방향성 포지션을 산출하는 규칙 기반 포트폴리오."
)

add_heading("1. 전략 내용", size=12)
add_bullet("팩터 판단: 각 팩터는 연속적인 z-score 시그널로 산출되어 자산군 내에서 합산됨.")
add_subbullet("Trend(추세): 6~12개월 다중 기간 시계열 모멘텀 → 추세추종(방향성 유지)")
add_subbullet("Value(가치): 2년 이동평균 대비 가격 z-score의 역방향 → 평균회귀")
add_subbullet("Carry(캐리): 국가별 국채 수익률 레벨의 횡단면 비교 → 고금리 롱 / 저금리 숏")
add_subbullet("Policy(정책): 정책금리(2Y/3Y) 변화 모멘텀 → 완화/긴축 사이클 방향 포착")
add_bullet(
    "포지션 산출: 팩터 결합 시그널 → 인버스 변동성 사이징(자산별 리스크 균등화) → "
    "포트폴리오 변동성 타겟팅(연 10%) → 최종 포지션."
)

add_heading("2. 리스크 관리", size=12)
add_bullet(
    "Book Stop(북 손절): 그림자(가상) 손익 기준 드로다운 4%p 초과 시 포지션 50% 축소, "
    "8%p 초과 시 전량 청산 후 회복 시 자동 재진입."
)
add_bullet(
    "xs-Reversion 서브북: 최근 10일 국채 횡단면 되돌림을 시장중립으로 상시 페이딩(북 스탑과 별도 운용). "
    "⚠ 거래비용 민감 — 왕복 1bp 초과 시 순수익 소멸, 패시브 집행 전제."
)
add_bullet("변동성 목표제: 시장 국면과 무관하게 리스크 예산(연 10%)을 항상 일정 유지.")

# ── [전략 2] FX 퀀트 전략 공장 ──────────────────────────────────────
add_heading("[전략 2] FX 퀀트 전략 공장 (Discovery Factory)")
add_body(
    "FX·금리 선물 유니버스 전반에서 통계적으로 유의한 구조적 크로스에셋 시그널만 "
    "자동으로 탐색·검증·저장하여 실거래에 투입하는 팩터 발굴 엔진."
)

add_heading("1. 전략 내용 (현재 라이브: Alpha1 + Alpha4)", size=12)
add_bullet("Alpha1 (횡단면 팩터): 자산군 내 모멘텀·캐리 순위 상위/하위 종목 롱숏.")
add_bullet(
    "Alpha4 (수익률 기반): 커브 슬로프(10Y-2Y)·정책금리 모멘텀·실질금리차 기반 "
    "FX 방향성 등 금리 데이터 활용 캐리/밸류 시그널."
)
add_bullet(
    "탐색은 하되 현재 라이브 제외: Advanced(필터드/멀티TF 모멘텀·상대강도 랭크), "
    "Alpha2(커브·금리차→FX)는 2016~2026 전/후반 포워드 샤프가 불안정/음(-)해 "
    "실거래 셀에서 빠짐(fx-advanced -0.47/-0.16, alpha2 -0.51/+0.24)."
)
add_bullet(
    "나이브 단일자산 기술지표(모멘텀·평균회귀)는 데이터마이닝성이 높아 "
    "2022~2026년 백테스트에서 순기여 無/음(-)으로 확인되어 탐색 단계부터 전면 제외."
)

add_heading("2. 리스크 관리", size=12)
add_bullet(
    "이중 아웃오브샘플 검증: 자산군×카테고리 조합별 2016~2026년 전/후반 분할, "
    "양쪽 구간 모두 양(+) 샤프인 조합만 실거래 활성화."
)
add_bullet("거래비용 반영: 왕복 2bp를 탐색·검증 단계에 상시 반영해 비용을 못 이기는 전략을 사전 배제.")
add_bullet("룩어헤드 차단: 시그널 1거래일 지연(signal_lag=1)으로 동일 봉 내 정보 선반영을 원천 차단.")
add_bullet("최소 자산수 하한 등 프레임워크 레벨 필터로 소수 자산 과최적화 방지.")

doc.save(r"D:\김선도\Python\only_quant\docs\전략공장_요약.docx")
print("saved")
