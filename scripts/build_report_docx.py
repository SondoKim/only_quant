"""Build a concise MS Word brief: 1-line summary + Q&A scripts."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUT_PATH = Path(r"D:/김선도/Python/only_quant/보고자료_운용전략_AI접목_v3.docx")

KOR_FONT = "맑은 고딕"
NAVY = RGBColor(0x1F, 0x3A, 0x68)
ACCENT = RGBColor(0xC0, 0x39, 0x2B)
GREY = RGBColor(0x55, 0x55, 0x55)
LIGHT_BG = "F2F4F8"


def set_run_font(run, name=KOR_FONT, size=11, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        from docx.oxml import OxmlElement

        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), name)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)


def add_para(doc, text="", size=11, bold=False, color=None, align=None,
             space_before=0, space_after=4):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = 1.45
    if text:
        run = p.add_run(text)
        set_run_font(run, size=size, bold=bold, color=color)
    return p


def add_section_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run_font(run, size=14, bold=True, color=NAVY)

    from docx.oxml import OxmlElement

    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "1F3A68")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def add_callout(doc, text, color=NAVY, size=12):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.right_indent = Cm(0.3)
    p.paragraph_format.line_spacing = 1.4
    run = p.add_run(text)
    set_run_font(run, size=size, bold=True, color=color)

    from docx.oxml import OxmlElement

    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), LIGHT_BG)
    pPr.append(shd)
    pBdr = OxmlElement("w:pBdr")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "6")
        b.set(qn("w:space"), "6")
        b.set(qn("w:color"), "1F3A68")
        pBdr.append(b)
    pPr.append(pBdr)
    return p


def add_qa(doc, idx, question, answer):
    # Q
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.4
    q_label = p.add_run(f"Q{idx}.  ")
    set_run_font(q_label, size=12, bold=True, color=ACCENT)
    q_text = p.add_run(question)
    set_run_font(q_text, size=12, bold=True, color=NAVY)

    # A
    a = doc.add_paragraph()
    a.paragraph_format.left_indent = Cm(0.7)
    a.paragraph_format.space_after = Pt(2)
    a.paragraph_format.line_spacing = 1.5
    a_label = a.add_run("A.  ")
    set_run_font(a_label, size=11, bold=True, color=GREY)
    a_text = a.add_run(answer)
    set_run_font(a_text, size=11)


def build():
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.4)
        section.right_margin = Cm(2.4)

    style = doc.styles["Normal"]
    style.font.name = KOR_FONT
    style.font.size = Pt(11)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        from docx.oxml import OxmlElement

        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), KOR_FONT)
    rfonts.set(qn("w:ascii"), KOR_FONT)
    rfonts.set(qn("w:hAnsi"), KOR_FONT)

    # ===== Title =====
    add_para(doc, "현업 개선 사례 — 발표 보조자료", size=10.5, color=GREY,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_para(doc, "운용전략에 AI 모델 접목", size=22, bold=True, color=NAVY,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=14)

    # ===== 요약 =====
    add_section_heading(doc, "요약")
    add_callout(doc,
                "채권 매니저의 매매 논리(모멘텀·평균회귀·캐리·금리 스프레드 등)를 룰로 코드화하고, "
                "AI 에이전트가 글로벌 8개국 금리·5개 통화 자산에 대해 수만 개 전략 조합을 "
                "자동 백테스트·검증한 뒤, 상관관계 필터와 성과 가중을 거쳐 매일 아침 "
                "자산별 매매 신호를 산출하는 시스템.",
                size=12)

    # ===== 예상 질문 스크립트 =====
    add_section_heading(doc, "예상 질문 & 답변 스크립트")

    add_qa(doc, 1,
           "AI 도입 전·후 가장 큰 변화는 무엇입니까?",
           "매매 아이디어 검증이 수일~수주에서 분 단위로 단축되었고, "
           "한 매니저가 다루던 수십 개 전략 대신 시스템이 수만 개를 동시에 검증합니다. "
           "그 결과 운용역의 시간을 단순 검증 작업에서 전략 설계·리스크 판단 같은 "
           "고차원 업무로 재배치할 수 있게 되었습니다.")

    add_qa(doc, 2,
           "AI가 매니저 판단을 대체하는 것 아닙니까? 블랙박스 리스크는 어떻게 관리합니까?",
           "저희가 도입한 것은 '블랙박스 예측 AI'가 아니라 '룰 기반 화이트박스 AI'입니다. "
           "매니저의 매매 논리(모멘텀·평균회귀·캐리·스프레드 등)를 사람이 읽을 수 있는 룰로 "
           "코드화하고, AI는 그중 어떤 룰·파라미터가 현 시장에 유효한지를 통계적으로 "
           "선별하는 역할만 합니다. 모든 진입·청산 조건이 추적·재현 가능하므로 "
           "리스크관리·컴플라이언스 대응에 즉시 활용할 수 있습니다.")

    add_qa(doc, 3,
           "수만 개 전략 중 우연히 잘 나온 전략이 채택될 위험(과적합)은 어떻게 차단합니까?",
           "세 가지 장치로 구조적으로 통제합니다. "
           "첫째, Walk-Forward 방식으로 매월 그 시점까지의 데이터만으로 전략을 발굴해 "
           "미래 데이터 누수를 차단합니다. "
           "둘째, 수익성 외에도 Sortino·거래횟수 등 복수 임계값을 통과한 전략만 채택합니다. "
           "셋째, 상관관계 필터로 서로 다른 알파 소스만 골라 우연성의 영향을 최소화하고 "
           "분산 효과를 확보합니다.")

    doc.save(OUT_PATH)
    print(f"Saved: {OUT_PATH}")


if __name__ == "__main__":
    build()
